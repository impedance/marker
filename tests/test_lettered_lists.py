import zipfile
from pathlib import Path

import pytest
from docx import Document

from core.adapters.document_parser import parse_document
from core.model.internal_doc import InternalDoc, ListBlock, ListItem, Paragraph, Text
from core.render.markdown_renderer import render_markdown


def test_render_lettered_list():
    """
    Tests that a ListBlock with list_style 'lowerLetter' is rendered
    into the custom ::letter-list format.
    """
    # 1. Create an InternalDoc with a lettered list
    doc = InternalDoc(
        blocks=[
            ListBlock(
                level=0,
                list_style="lowerLetter",  # The new attribute to indicate list type
                items=[
                    ListItem(
                        level=0,
                        blocks=[Paragraph(inlines=[Text(content="первый способ")])],
                    ),
                    ListItem(
                        level=0,
                        blocks=[Paragraph(inlines=[Text(content="второй способ")])],
                    ),
                ],
            )
        ]
    )

    # 2. Render the document to Markdown
    markdown_output = render_markdown(doc, {})

    # 3. Assert the output is in the correct format
    expected_output = (
        "::letter-list\n"
        "- первый способ\n"
        "- второй способ\n"
        "::\n"
    )

    assert markdown_output.strip() == expected_output.strip()


def test_parse_lettered_list_from_docx(tmp_path: Path) -> None:
    """Lists with letter numbering should render using ::letter-list format."""

    doc = Document()
    doc.add_paragraph("первый способ", style="List Number")
    doc.add_paragraph("второй способ", style="List Number")

    docx_path = tmp_path / "lettered.docx"
    doc.save(docx_path)

    with zipfile.ZipFile(docx_path, "r") as existing_zip:
        file_contents = {name: existing_zip.read(name) for name in existing_zip.namelist()}

    numbering_xml = file_contents.get("word/numbering.xml", b"")
    assert numbering_xml, "numbering.xml must be present in DOCX archive"

    if b'val="lowerLetter"' not in numbering_xml:
        updated_numbering = numbering_xml.replace(b'val="decimal"', b'val="lowerLetter"')
        assert updated_numbering != numbering_xml, "List numbering format replacement should occur"
        file_contents["word/numbering.xml"] = updated_numbering

        with zipfile.ZipFile(docx_path, "w") as updated_zip:
            for name, content in file_contents.items():
                updated_zip.writestr(name, content)

    internal_doc, _ = parse_document(str(docx_path))

    list_block = next(block for block in internal_doc.blocks if isinstance(block, ListBlock))
    assert list_block.list_style == "lowerLetter"

    markdown_output = render_markdown(internal_doc, {})

    expected = (
        "::letter-list\n"
        "- первый способ\n"
        "- второй способ\n"
        "::\n"
    )

    assert expected.strip() in markdown_output.strip()
