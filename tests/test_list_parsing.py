from pathlib import Path

from docx import Document

from core.adapters.document_parser import parse_document
from core.model.internal_doc import (
    InternalDoc,
    Paragraph,
    Text,
    Bold,
    ListBlock,
    ListItem,
)
from core.render.markdown_renderer import render_markdown


def test_bullet_list_parsed(tmp_path: Path) -> None:
    """Ensure bullet lists are parsed with inline formatting."""
    doc = Document()
    first = doc.add_paragraph(style="List Bullet")
    first_run = first.add_run("Standard")
    first_run.bold = True
    first.add_run(" item")
    doc.add_paragraph("Second", style="List Bullet")
    path = tmp_path / "list.docx"
    doc.save(path)

    internal_doc, _ = parse_document(str(path))
    assert internal_doc.blocks
    list_block = internal_doc.blocks[0]
    assert isinstance(list_block, ListBlock)
    assert len(list_block.items) == 2

    first_item = list_block.items[0]
    assert first_item.blocks
    first_paragraph = first_item.blocks[0]
    assert isinstance(first_paragraph, Paragraph)
    assert isinstance(first_paragraph.inlines[0], Bold)
    assert first_paragraph.inlines[0].content == "Standard"

    markdown = render_markdown(internal_doc, {})
    assert markdown.splitlines()[:2] == ["- **Standard** item", "- Second"]


def test_nested_list_parsed(tmp_path: Path) -> None:
    """Lists with indentation are captured as nested structures."""
    doc = Document()
    doc.add_paragraph("Parent", style="List Bullet")
    doc.add_paragraph("Child", style="List Bullet 2")
    path = tmp_path / "nested.docx"
    doc.save(path)

    internal_doc, _ = parse_document(str(path))
    assert internal_doc.blocks
    list_block = internal_doc.blocks[0]
    assert isinstance(list_block, ListBlock)
    assert len(list_block.items) == 1

    nested_blocks = list_block.items[0].blocks
    assert len(nested_blocks) >= 2
    nested_list = next(
        block for block in nested_blocks if isinstance(block, ListBlock)
    )
    nested_paragraph = nested_list.items[0].blocks[0]
    assert isinstance(nested_paragraph, Paragraph)
    nested_text = "".join(inline.content for inline in nested_paragraph.inlines)
    assert nested_text.strip() == "Child"

    markdown = render_markdown(internal_doc, {})
    assert "- Parent" in markdown
    assert "  - Child" in markdown


def test_nested_list_rendering() -> None:
    """Renderer indents nested list items by level."""
    doc = InternalDoc(
        blocks=[
            ListBlock(
                ordered=False,
                items=[
                    ListItem(
                        blocks=[
                            Paragraph(inlines=[Text(content="Parent")]),
                            ListBlock(
                                ordered=False,
                                items=[
                                    ListItem(
                                        blocks=[
                                            Paragraph(
                                                inlines=[Text(content="Nested item")]
                                            )
                                        ]
                                    )
                                ],
                            ),
                        ]
                    )
                ],
            )
        ]
    )

    markdown = render_markdown(doc, {})
    assert markdown == "- Parent\n  - Nested item"


def test_list_renderer_escapes_blockquote_marker() -> None:
    """Leading blockquote markers in list items are escaped."""
    doc = InternalDoc(
        blocks=[
            ListBlock(
                ordered=False,
                items=[
                    ListItem(
                        blocks=[
                            Paragraph(
                                inlines=[
                                    Text(
                                        content=">& файл – переадресовывает сообщения"
                                    )
                                ]
                            )
                        ]
                    )
                ],
            )
        ]
    )

    markdown = render_markdown(doc, {})
    assert markdown.startswith("- \\>& файл – переадресовывает сообщения")
